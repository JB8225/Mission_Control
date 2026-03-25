#!/usr/bin/env python3
"""Build 4 Word documents for One Buck Capital / GFS deals."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Brand colors
NAVY = RGBColor(0x1B, 0x2A, 0x3D)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
OFFWHITE = RGBColor(0xF5, 0xF0, 0xE8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)

def set_cell_bg(cell, hex_color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), str(val.get('sz', 4)))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=None, size=None, center=False, bold=True, font='Georgia'):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    elif level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    elif level == 3:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    return p

def add_body(doc, text, size=11, color=None, bold=False, italic=False, center=False, font='Calibri'):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, size=11, color=None, font='Calibri', bold=False):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_divider(doc, color='1B2A3D'):
    """Add a horizontal rule-like paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    return p

def set_page_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def add_navy_header_bar(doc, title, subtitle=None):
    """Add a full-width navy header table."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, '1B2A3D')
    cell.width = Inches(6.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(20)
    run.font.color.rgb = WHITE
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(8)
        run2 = p2.add_run(subtitle)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(11)
        run2.font.color.rgb = GOLD
    else:
        p.paragraph_format.space_after = Pt(10)
    doc.add_paragraph()
    return table


# ============================================================
# DOCUMENT 1: NNN Broker List
# ============================================================
def build_nnn_broker_list():
    doc = Document()
    set_page_margins(doc, 1.0, 1.0, 1.1, 1.1)

    add_navy_header_bar(doc, 'ONE BUCK CAPITAL', 'NNN Broker Contact Directory — March 2026')

    add_heading(doc, 'Top NNN Brokers — Starbucks/QSR Exits & 1031 Exchange Specialists', 2, NAVY)
    add_body(doc, 'Target States: Missouri | Texas | Colorado     Focus: NNN single-tenant exits, 1031 exchange expertise', italic=True, color=MID_GRAY)
    add_divider(doc)

    brokers = [
        {
            'name': '1. Westwood Net Lease Advisors',
            'flag': '⭐ RECOMMENDED',
            'hq': '1401 S. Brentwood Blvd, Suite 650, Brentwood, MO 63144',
            'contacts': [
                ('President / Principal', 'Chris Schellin', '314-563-2208', 'cschellin@westwoodnetlease.com'),
                ('VP / Principal', 'Jason Simon', '314-563-2206', 'jsimon@westwoodnetlease.com'),
                ('VP / Principal', 'Vince Vatterott', '314-563-2204', 'vvatterott@westwoodnetlease.com'),
                ('VP / Principal', 'Mike Kocur', '314-266-2654', 'mkocur@westwoodnetlease.com'),
                ('General Line', '—', '314-997-5227', 'info@westwoodnetlease.com'),
            ],
            'bullets': [
                'Missouri-based boutique NNN shop — they live and breathe this asset class. Deep 1031 exchange expertise is core to their business model.',
                'Explicitly markets Starbucks NNN properties across Texas and Colorado — strong cross-state coverage for seller exits and 1031 replacement sourcing.',
                'No-obligation consultation offered. Easy, low-friction first call. Strong seller-side and buyer-side representation.',
            ],
        },
        {
            'name': '2. Secure Net Lease',
            'flag': '⭐ RECOMMENDED',
            'hq': '10000 N. Central Expressway #200, Dallas, TX 75231',
            'contacts': [
                ('Founding Partner', 'Bob Moorhead', '214-915-8890', '—'),
                ('Founding Partner', 'Joe Caputo', '424-220-6432', '—'),
                ('EVP', 'Edward Benton', '713-263-3981', '—'),
                ('EVP', 'Anthony Pucciarello', '214-915-8896', '—'),
                ('General Line', '—', '214-522-7200', 'info_dallas@securenetlease.com'),
            ],
            'bullets': [
                'Documented Starbucks NNN closings in both Texas (Carrollton) and Colorado (Greeley) — rare cross-state coverage matching our target markets exactly.',
                'Dallas-based with strong institutional relationships with Starbucks licensees and QSR operators looking to exit. Deep Texas market intel.',
                'Experienced in assisting 1031 exchange buyers find replacement NNN assets — handles both sides of the transaction.',
            ],
        },
        {
            'name': '3. NNN Retail Advisors',
            'flag': None,
            'hq': '18208 Preston Road, Suite D9-278, Dallas, TX 75252',
            'contacts': [
                ('Partner', 'Gavin M. Kam', '972-375-3437', 'gavin@nnnretailadvisors.com'),
                ('Partner', 'Brad F. Kam', '972-375-3438', 'brad@nnnretailadvisors.com'),
            ],
            'bullets': [
                'Closed 900+ transactions, $3B+ in volume since 2002. Pure-play boutique NNN firm — not a generalist shop. Deep track record.',
                'QSR and Starbucks tenant expertise with strong Texas market coverage. 1031 exchange buyer/seller representation is their bread and butter.',
                'Dallas-based boutique: more accessible and deal-focused than nationals. Better responsiveness and personal service.',
            ],
        },
        {
            'name': '4. Blue West Capital',
            'flag': None,
            'hq': '650 S Cherry St, Suite 920, Denver, CO 80246',
            'contacts': [
                ('Broker — Starbucks', 'Tom Ethington', '720-232-9530', '—'),
                ('Broker — Starbucks', 'Robert Edwards', '720-966-1630', '—'),
                ('General', '—', '—', 'info@BlueWestCapital.com'),
            ],
            'bullets': [
                'Directly completed sale of Starbucks NNN in Centennial, CO — documented closing in our primary target state. Colorado-native with deep local retail relationships.',
                'Tom Ethington and Robert Edwards are the go-to brokers for Starbucks/QSR in Colorado. Active deal flow in Denver metro, Colorado Springs, and Front Range.',
                '1031 exchange specialists on staff. Local expertise translates to faster, better-priced inventory access.',
            ],
        },
        {
            'name': '5. Fountainhead Commercial',
            'flag': None,
            'hq': 'Denver, CO',
            'contacts': [
                ('Founder / CCIM', 'Lowrey Burnett', '720-837-9407', 'LBurnett@FountainheadCommercial.com'),
            ],
            'bullets': [
                'CCIM-designated 1031 Exchange specialist — certified expert, not a generalist. 450+ commercial transactions, nearly 20 years in Colorado commercial real estate.',
                'Licensed to execute 1031 exchanges across all 50 states — critical if the replacement property is out-of-state. Rare capability in a boutique firm.',
                'Smaller shop means more hands-on, personal service. Ideal for a first 1031 transaction where hand-holding and guidance add real value.',
            ],
        },
    ]

    for broker in brokers:
        doc.add_paragraph()
        # Broker name heading
        p = doc.add_paragraph()
        run = p.add_run(broker['name'])
        run.bold = True
        run.font.name = 'Georgia'
        run.font.size = Pt(14)
        run.font.color.rgb = NAVY
        if broker['flag']:
            run2 = p.add_run(f'  {broker["flag"]}')
            run2.bold = True
            run2.font.name = 'Calibri'
            run2.font.size = Pt(10)
            run2.font.color.rgb = GOLD

        # Address
        p_addr = doc.add_paragraph()
        run_addr = p_addr.add_run(f'📍  {broker["hq"]}')
        run_addr.font.name = 'Calibri'
        run_addr.font.size = Pt(10)
        run_addr.font.color.rgb = MID_GRAY
        p_addr.paragraph_format.space_after = Pt(4)

        # Contacts table
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        # Header row
        hdr = tbl.rows[0]
        headers = ['Title', 'Name', 'Phone', 'Email']
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            set_cell_bg(cell, '1B2A3D')
            p2 = cell.paragraphs[0]
            run2 = p2.add_run(h)
            run2.bold = True
            run2.font.name = 'Calibri'
            run2.font.size = Pt(9)
            run2.font.color.rgb = WHITE

        for contact in broker['contacts']:
            row = tbl.add_row()
            for i, val in enumerate(contact):
                cell = row.cells[i]
                p2 = cell.paragraphs[0]
                run2 = p2.add_run(val)
                run2.font.name = 'Calibri'
                run2.font.size = Pt(9)
                run2.font.color.rgb = DARK_GRAY
                if i % 2 == 0:
                    set_cell_bg(cell, 'F5F0E8')
                else:
                    set_cell_bg(cell, 'FFFFFF')

        # Set col widths
        for i, width in enumerate([1.5, 1.4, 1.2, 2.4]):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(width)

        doc.add_paragraph()
        # Bullets
        add_body(doc, 'Why They\'re a Good Fit:', bold=True, size=10, color=NAVY)
        for bullet in broker['bullets']:
            add_bullet(doc, bullet, size=10)

        add_divider(doc)

    # Honorable mentions
    doc.add_paragraph()
    add_heading(doc, 'Honorable Mentions', 3, NAVY)

    hon_tbl = doc.add_table(rows=1, cols=4)
    hon_tbl.style = 'Table Grid'
    for i, h in enumerate(['Company', 'State', 'Specialty', 'Note']):
        cell = hon_tbl.rows[0].cells[i]
        set_cell_bg(cell, '1B2A3D')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    mentions = [
        ('Hilliker Corp', 'MO (St. Louis)', 'NNN + 1031', 'Partners with Westwood; good secondary contact'),
        ('Marcus & Millichap St. Louis', 'MO', 'NNN Retail', 'National resources, good for Starbucks inventory access'),
        ('Quantum Real Estate Advisors', 'CO', 'NNN (Starbucks Lakewood)', 'Smaller boutique, active CO listings'),
    ]
    for i, row_data in enumerate(mentions):
        row = hon_tbl.add_row()
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            run = cell.paragraphs[0].add_run(val)
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            if i % 2 == 0:
                set_cell_bg(cell, 'F5F0E8')

    # Recommended first call callout
    doc.add_paragraph()
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_bg(cell, 'C9A84C')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('★  RECOMMENDED FIRST CALLS  ★')
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(2)
    for line in [
        'Seller Exits (Starbucks/QSR owner selling):\n  → Westwood Net Lease (MO) + Secure Net Lease (TX/CO) — both sides covered.',
        '1031 Exchange Buyers (finding replacement NNN):\n  → NNN Retail Advisors (TX) + Fountainhead Commercial (CO).',
        'Colorado-Specific Deals:\n  → Blue West Capital — proven local Starbucks track record.',
    ]:
        p3 = cell.add_paragraph()
        p3.paragraph_format.left_indent = Inches(0.2)
        p3.paragraph_format.space_before = Pt(2)
        p3.paragraph_format.space_after = Pt(2)
        run3 = p3.add_run(line)
        run3.font.name = 'Calibri'
        run3.font.size = Pt(10)
        run3.font.color.rgb = NAVY

    p_end = cell.add_paragraph()
    p_end.paragraph_format.space_after = Pt(8)

    # Footer
    doc.add_paragraph()
    add_body(doc, 'One Buck Capital — Confidential | Contact info current as of March 2026. Verify before outreach.',
             size=8, color=MID_GRAY, italic=True, center=True)

    doc.save('/home/node/.openclaw/workspace/NNN_Broker_List.docx')
    print('✓ NNN_Broker_List.docx saved')


# ============================================================
# DOCUMENT 2: GFS Woody One-Pager
# ============================================================
def build_gfs_woody_one_pager():
    doc = Document()
    set_page_margins(doc, 0.75, 0.75, 1.0, 1.0)

    # Big navy header
    add_navy_header_bar(doc, 'THE CHOICE IN FRONT OF YOU', 'GFS Fire Pros — What This Deal Actually Means for Your Family')

    # Big bold headline
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Two Paths. One Real Number.')
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY

    doc.add_paragraph()

    # Comparison table
    comp_tbl = doc.add_table(rows=6, cols=3)
    comp_tbl.style = 'Table Grid'
    comp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['', 'Cash Offer', 'Viking 90 Partnership']
    row0 = comp_tbl.rows[0]
    for i, h in enumerate(headers):
        cell = row0.cells[i]
        set_cell_bg(cell, '1B2A3D')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE if h else GOLD
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    comp_data = [
        ('Headline Price', '$14,000,000', '$19,000,000+'),
        ('Broker Fees', '−$700,000', '$0'),
        ('Capital Gains + Taxes', '−$3,540,000', 'Spread over 5+ years'),
        ('What You Actually Keep', '$9,760,000', '$14,000,000+'),
    ]

    for row_idx, (label, cash, viking) in enumerate(comp_data):
        row = comp_tbl.rows[row_idx + 1]
        is_last = row_idx == len(comp_data) - 1

        cell0 = row.cells[0]
        cell0.paragraphs[0].add_run(label).font.name = 'Calibri'
        cell0.paragraphs[0].runs[0].font.size = Pt(10)
        cell0.paragraphs[0].runs[0].bold = is_last
        if is_last:
            set_cell_bg(cell0, 'F5F0E8')
        else:
            set_cell_bg(cell0, 'FFFFFF')

        cell1 = row.cells[1]
        r1 = cell1.paragraphs[0].add_run(cash)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(10)
        r1.bold = is_last
        if is_last:
            set_cell_bg(cell1, 'F5F0E8')
        else:
            set_cell_bg(cell1, 'FFFFFF')
        cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell2 = row.cells[2]
        r2 = cell2.paragraphs[0].add_run(viking)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(10)
        r2.bold = is_last
        if is_last:
            r2.font.color.rgb = NAVY
            set_cell_bg(cell2, 'C9A84C')
        else:
            set_cell_bg(cell2, 'FFFFFF')
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Col widths
    for i, w in enumerate([2.2, 2.0, 2.3]):
        for cell in comp_tbl.columns[i].cells:
            cell.width = Inches(w)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('The cash offer sounds bigger. It isn\'t.')
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(13)
    run.font.color.rgb = GOLD

    add_divider(doc)

    # How the money flows
    add_heading(doc, 'How the Money Flows to You', 2, NAVY)

    flow_tbl = doc.add_table(rows=4, cols=2)
    flow_tbl.style = 'Table Grid'

    flow_data = [
        ('Day 1', '$8,200,000 cash at closing\n($5.7M for the business + $2.5M for the real estate)'),
        ('Every Month\nfor 5 Years', '$38,666/month — seller note, principal + interest, like clockwork\n$2,319,935 total over 60 months'),
        ('Every Quarter', '18% of company profits distributed as an equity holder\nEstimated $752,000–$1,103,000 over 5 years'),
        ('Year 5', 'Your 18% equity stake bought out at fair market value\nEstimated $2,900,000–$5,000,000 depending on growth'),
    ]

    for row_idx, (timeframe, detail) in enumerate(flow_data):
        row = flow_tbl.rows[row_idx]
        cell0 = row.cells[0]
        set_cell_bg(cell0, '1B2A3D')
        r = cell0.paragraphs[0].add_run(timeframe)
        r.bold = True
        r.font.name = 'Georgia'
        r.font.size = Pt(10)
        r.font.color.rgb = GOLD
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        cell1 = row.cells[1]
        if row_idx % 2 == 0:
            set_cell_bg(cell1, 'F5F0E8')
        r2 = cell1.paragraphs[0].add_run(detail)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(10)
        cell1.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell1.paragraphs[0].paragraph_format.space_after = Pt(3)

    for cell in flow_tbl.columns[0].cells:
        cell.width = Inches(1.4)
    for cell in flow_tbl.columns[1].cells:
        cell.width = Inches(5.1)

    doc.add_paragraph()

    # Total callout
    total_tbl = doc.add_table(rows=1, cols=1)
    total_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = total_tbl.cell(0, 0)
    set_cell_bg(cell, 'C9A84C')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('YOUR TOTAL:  $14,000,000 to $16,400,000')
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run('Compared to $9,760,000 from a cash sale — after taxes and fees.')
    run2.font.name = 'Calibri'
    run2.font.size = Pt(10)
    run2.font.color.rgb = NAVY

    add_divider(doc)

    # Two columns section - What you're not giving up + Why $19M is the floor
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('What You\'re Not Giving Up')
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY

    keeps = [
        'GFS keeps its name — 40 years of brand equity preserved.',
        'Your team stays — with equity incentives to grow.',
        'You remain an 18% equity holder with advisory rights.',
        'Carly, Chad, and the grandchildren benefit from continued growth.',
        'A voice in the company\'s future — you\'re not walking away.',
    ]
    for k in keeps:
        add_bullet(doc, k, size=10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Why $19M Is the Floor, Not the Ceiling')
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY

    add_body(doc, 'GFS grew 34% last year. DFW is one of the fastest-growing markets in the country. We\'re investing in sales, service, and expansion — and your 18% rides with every dollar.',
             size=10)
    add_body(doc, 'If GFS reaches $50M in revenue over the next 7–10 years, your 18% stake is worth $9–12 million at exit.',
             size=10, bold=True, color=NAVY)
    add_body(doc, 'The cash buyer ends the story.  This deal starts a new chapter — and you\'re in it.',
             size=10, italic=True, color=MID_GRAY)

    # Footer
    add_divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Joel Bouck — Viking 90 Group  |  jb@onebuckcapital.com  |  Confidential')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = MID_GRAY

    doc.save('/home/node/.openclaw/workspace/GFS_Woody_One_Pager.docx')
    print('✓ GFS_Woody_One_Pager.docx saved')


# ============================================================
# DOCUMENT 3: GFS SafeGuard Program
# ============================================================
def build_gfs_safeguard():
    doc = Document()
    set_page_margins(doc, 1.0, 1.0, 1.1, 1.1)

    add_navy_header_bar(doc, 'GFS SAFEGUARD™', 'Safety as a Service — Program Overview & Business Case')

    # Subtitle
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('GFS Fire Pros  |  Viking 90 Group  |  CONFIDENTIAL')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = MID_GRAY

    add_divider(doc)

    # Program overview
    add_heading(doc, 'PART 1: THE OPPORTUNITY', 2, NAVY)
    add_heading(doc, 'The Old Model vs. The New Model', 3, DARK_GRAY)

    model_tbl = doc.add_table(rows=5, cols=2)
    model_tbl.style = 'Table Grid'

    for i, h in enumerate(['The Old Model (Today)', 'The New Model (GFS SafeGuard)']):
        cell = model_tbl.rows[0].cells[i]
        set_cell_bg(cell, '1B2A3D' if i == 0 else 'C9A84C')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE if i == 0 else NAVY

    old_rows = [
        'Install a system → get paid → hope they call back',
        'Inspection is a one-time transaction',
        'Revenue is lumpy and unpredictable',
        'Customer calls a competitor when something breaks',
    ]
    new_rows = [
        'Install a system → lock in monthly revenue forever',
        'Every building on a subscription contract',
        'Revenue is predictable and compounds every month',
        'You\'re the first call because you never left',
    ]

    for i, (old, new) in enumerate(zip(old_rows, new_rows)):
        row = model_tbl.rows[i + 1]
        for j, text in enumerate([old, new]):
            cell = row.cells[j]
            if i % 2 == 0:
                set_cell_bg(cell, 'F5F0E8')
            run = cell.paragraphs[0].add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    doc.add_paragraph()

    # Proof point
    proof_tbl = doc.add_table(rows=1, cols=1)
    cell = proof_tbl.cell(0, 0)
    set_cell_bg(cell, '1B2A3D')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run('The Proof: How Pye-Barker Did It')
    run.bold = True
    run.font.name = 'Georgia'
    run.font.size = Pt(12)
    run.font.color.rgb = GOLD
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Inches(0.1)
    run2 = p2.add_run(
        'Pye-Barker started as a small regional fire protection company. They built a subscription/service model and completed 57 acquisitions in 2025 alone. '
        'Revenue: $2 billion. Monthly Recurring Revenue: $16M+ and growing. They didn\'t get there by being the best installer — '
        'they got there by owning the relationship after the install.\n\n'
        'GFS already has the relationships. SafeGuard turns them into monthly revenue.'
    )
    run2.font.name = 'Calibri'
    run2.font.size = Pt(10)
    run2.font.color.rgb = OFFWHITE

    add_divider(doc)

    # Three tiers
    add_heading(doc, 'PART 2: THE PROGRAM — THREE TIERS', 2, NAVY)

    tiers = [
        {
            'name': 'BRONZE — "Comply"',
            'color': 'B87333',
            'who': 'Small commercial buildings, single-tenant spaces, cost-conscious owners',
            'price': '$175–$350/month per building',
            'includes': [
                'Annual fire alarm inspection (NFPA 72)',
                'Annual sprinkler inspection (NFPA 25)',
                'Annual fire extinguisher inspection',
                'Digital inspection report delivered same day',
                'Compliance certificate for insurance and AHJ',
                'Online customer portal — all records 24/7',
            ],
            'metrics': [
                ('Annual Value to Customer', '$900–$2,400'),
                ('GFS Delivery Cost', '~$400–$600/year'),
                ('Gross Margin', '60–70%'),
            ],
        },
        {
            'name': 'SILVER — "Protect"',
            'color': '708090',
            'who': 'Mid-size commercial, property managers, multi-tenant buildings',
            'price': '$400–$700/month per building',
            'includes': [
                'Everything in Bronze, PLUS:',
                '24/7 fire alarm monitoring (white-label central station)',
                'Semi-annual inspections (alarm + sprinkler)',
                'Deficiency tracking — repair quote within 24 hours',
                'Priority response (within 4 hours for service calls)',
                'Monthly system status email to building owner',
                'Dedicated GFS account manager',
            ],
            'metrics': [
                ('Annual Value to Customer', '$4,800–$8,400'),
                ('GFS Delivery Cost', '~$1,200–$2,000/year'),
                ('Gross Margin', '65–75%'),
            ],
        },
        {
            'name': 'GOLD — "Command"',
            'color': 'C9A84C',
            'who': 'Data centers, hospitals, large commercial, multi-location operators',
            'price': '$800–$1,500/month per building',
            'includes': [
                'Everything in Silver, PLUS:',
                'Quarterly inspections across all systems',
                'AI-powered predictive monitoring (IoT sensors)',
                'Annual system health report + 12-month maintenance schedule',
                'Same-day emergency response guarantee',
                'Dedicated technician assigned to each building',
                'Executive dashboard — real-time health and compliance status',
                'Insurance documentation package (audit-ready)',
            ],
            'metrics': [
                ('Annual Value to Customer', '$9,600–$18,000'),
                ('GFS Delivery Cost', '~$2,500–$4,000/year'),
                ('Gross Margin', '70–80%'),
            ],
        },
    ]

    for tier in tiers:
        doc.add_paragraph()
        # Tier header
        tier_hdr = doc.add_table(rows=1, cols=2)
        tier_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_l = tier_hdr.cell(0, 0)
        cell_r = tier_hdr.cell(0, 1)
        set_cell_bg(cell_l, tier['color'])
        set_cell_bg(cell_r, tier['color'])

        run_l = cell_l.paragraphs[0].add_run(tier['name'])
        run_l.bold = True
        run_l.font.name = 'Georgia'
        run_l.font.size = Pt(14)
        run_l.font.color.rgb = WHITE if tier['color'] != 'C9A84C' else NAVY
        cell_l.paragraphs[0].paragraph_format.space_before = Pt(6)
        cell_l.paragraphs[0].paragraph_format.space_after = Pt(6)

        run_r = cell_r.paragraphs[0].add_run(tier['price'])
        run_r.bold = True
        run_r.font.name = 'Georgia'
        run_r.font.size = Pt(14)
        run_r.font.color.rgb = WHITE if tier['color'] != 'C9A84C' else NAVY
        cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell_r.paragraphs[0].paragraph_format.space_before = Pt(6)
        cell_r.paragraphs[0].paragraph_format.space_after = Pt(6)

        cell_l.width = Inches(3.5)
        cell_r.width = Inches(3.0)

        # Who it's for
        p_who = doc.add_paragraph()
        run_who = p_who.add_run(f'Best for: {tier["who"]}')
        run_who.italic = True
        run_who.font.name = 'Calibri'
        run_who.font.size = Pt(10)
        run_who.font.color.rgb = MID_GRAY

        # Includes + metrics table
        content_tbl = doc.add_table(rows=1, cols=2)
        content_tbl.style = 'Table Grid'
        cell_inc = content_tbl.cell(0, 0)
        cell_met = content_tbl.cell(0, 1)
        set_cell_bg(cell_inc, 'FFFFFF')
        set_cell_bg(cell_met, 'F5F0E8')

        run_inc_hdr = cell_inc.paragraphs[0].add_run('What\'s Included')
        run_inc_hdr.bold = True
        run_inc_hdr.font.name = 'Calibri'
        run_inc_hdr.font.size = Pt(10)
        run_inc_hdr.font.color.rgb = NAVY

        for item in tier['includes']:
            p_item = cell_inc.add_paragraph()
            run_item = p_item.add_run(f'✓  {item}')
            run_item.font.name = 'Calibri'
            run_item.font.size = Pt(9)
            p_item.paragraph_format.space_before = Pt(1)
            p_item.paragraph_format.space_after = Pt(1)

        run_met_hdr = cell_met.paragraphs[0].add_run('Economics')
        run_met_hdr.bold = True
        run_met_hdr.font.name = 'Calibri'
        run_met_hdr.font.size = Pt(10)
        run_met_hdr.font.color.rgb = NAVY

        for label, val in tier['metrics']:
            p_m = cell_met.add_paragraph()
            run_m = p_m.add_run(f'{label}:\n')
            run_m.bold = True
            run_m.font.name = 'Calibri'
            run_m.font.size = Pt(9)
            run_m2 = p_m.add_run(val)
            run_m2.font.name = 'Calibri'
            run_m2.font.size = Pt(11)
            run_m2.bold = True
            run_m2.font.color.rgb = NAVY
            p_m.paragraph_format.space_before = Pt(4)

        cell_inc.width = Inches(4.0)
        cell_met.width = Inches(2.5)

    add_divider(doc)

    # How to launch
    add_heading(doc, 'PART 3: HOW TO LAUNCH — STEP BY STEP', 2, NAVY)

    steps = [
        ('Step 1 — Pick Your Monitoring Partner (Month 1–2)',
         'White-label monitoring from a wholesale partner. GFS SafeGuard brand on the outside; their infrastructure underneath.\n'
         'Wholesale cost to GFS: $8–$15/account/month.  Charge to customer: $50–$75/month.\n'
         'Recommended partners: Alarm.com · Monitoring America · COPS Monitoring · Affiliated Monitoring · Bold Group'),
        ('Step 2 — Pick Your Inspection Software (Month 1)',
         'Recommendation: InspectPoint (inspectpoint.com) — purpose-built for fire protection, AI-native, 15,000+ pros, 4.5M inspections processed.\n'
         'Saves 15–20 minutes per inspection. Customer portal included. Cost: ~$200–$400/month.'),
        ('Step 3 — Price Your Contracts (Month 2)',
         'Target gross margin: 65–75%. Calculate cost per building (labor + vehicle + monitoring + software overhead + 20% overhead allocation). Work backward to monthly price.'),
        ('Step 4 — The Conversion Play (Month 1–6)',
         'Pull every GFS job completed in the last 7 years. For each building with no service contract: assign to Sales Rep → call script → email follow-up within 2 hours → follow up in 3 business days.\n'
         'Conversion target: 25–35% of outreach in first 90 days. Close rate after proposal: 40–50%.'),
        ('Step 5 — New Install Policy (Permanent)',
         'Every new install leaves with a SafeGuard agreement signed at close. Close rate on new installs: 60–70% (they trust you, you\'re already there).'),
    ]

    for step_title, step_body in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(step_title)
        run.bold = True
        run.font.name = 'Georgia'
        run.font.size = Pt(11)
        run.font.color.rgb = NAVY
        add_body(doc, step_body, size=10)

    add_divider(doc)

    # Revenue projections
    add_heading(doc, 'PART 4: REVENUE PROJECTIONS — THREE SCENARIOS', 2, NAVY)

    scenarios = [
        {
            'name': 'SCENARIO 1: CONSERVATIVE',
            'subtitle': '10% organic growth, no acquisitions — "The Floor"',
            'color': '708090',
            'rows': [
                ('Month 6', '50', '$17,500', '$210,000', '~1%'),
                ('Month 12', '80', '$28,000', '$336,000', '~1.5%'),
                ('Month 24', '120', '$42,000', '$504,000', '~2%'),
                ('Month 36', '160', '$56,000', '$672,000', '~2.5%'),
            ],
            'note': 'Year 5 refi valuation: ~$28M (4.5–5x EBITDA)',
        },
        {
            'name': 'SCENARIO 2: BASE CASE',
            'subtitle': 'Focused execution + 1 tuck-in acquisition — "The Target"',
            'color': '1B2A3D',
            'rows': [
                ('Month 6', '75', '$26,250', '$315,000', '~1.5%'),
                ('Month 12', '150', '$52,500', '$630,000', '~3%'),
                ('Month 24', '300', '$105,000', '$1,260,000', '~5%'),
                ('Month 36', '500', '$175,000', '$2,100,000', '~8%'),
            ],
            'note': 'Company revenue ~$30M by Year 3  |  Year 5 refi valuation: ~$35–40M (5.5–6x)',
        },
        {
            'name': 'SCENARIO 3: REACH',
            'subtitle': 'Aggressive execution + 2 acquisitions — "Step on It"',
            'color': 'B87333',
            'rows': [
                ('Month 6', '100', '$35,000', '$420,000', '~2%'),
                ('Month 12', '250', '$87,500', '$1,050,000', '~5%'),
                ('Month 24', '500', '$175,000', '$2,100,000', '~8%'),
                ('Month 36', '900', '$315,000', '$3,780,000', '~13%'),
            ],
            'note': 'Company revenue ~$38–42M by Year 3  |  Year 5 refi valuation: ~$50–60M (6.5–7x)',
        },
    ]

    col_headers = ['Period', 'Accounts', 'Monthly ITM Revenue', 'Annual ITM Revenue', '% of Total Revenue']

    for scenario in scenarios:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(scenario['name'])
        run.bold = True
        run.font.name = 'Georgia'
        run.font.size = Pt(12)
        run.font.color.rgb = NAVY
        run2 = p.add_run(f'  —  {scenario["subtitle"]}')
        run2.font.name = 'Calibri'
        run2.font.size = Pt(10)
        run2.italic = True
        run2.font.color.rgb = MID_GRAY

        tbl = doc.add_table(rows=len(scenario['rows']) + 1, cols=5)
        tbl.style = 'Table Grid'

        for i, h in enumerate(col_headers):
            cell = tbl.rows[0].cells[i]
            set_cell_bg(cell, scenario['color'])
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE

        for r_idx, row_data in enumerate(scenario['rows']):
            row = tbl.rows[r_idx + 1]
            for c_idx, val in enumerate(row_data):
                cell = row.cells[c_idx]
                if r_idx % 2 == 0:
                    set_cell_bg(cell, 'F5F0E8')
                run = cell.paragraphs[0].add_run(val)
                run.font.name = 'Calibri'
                run.font.size = Pt(9)

        p_note = doc.add_paragraph()
        p_note.paragraph_format.space_before = Pt(2)
        run_note = p_note.add_run(f'★  {scenario["note"]}')
        run_note.italic = True
        run_note.font.name = 'Calibri'
        run_note.font.size = Pt(9)
        run_note.font.color.rgb = GOLD

    add_divider(doc)

    # Valuation impact
    add_heading(doc, 'PART 5: THE VALUATION IMPACT', 2, NAVY)
    add_body(doc, 'Why Recurring Revenue Changes the Multiple — Same Earnings. Different Story. Different Price.', bold=True, size=11, color=DARK_GRAY)

    val_tbl = doc.add_table(rows=5, cols=4)
    val_tbl.style = 'Table Grid'
    val_headers = ['Revenue Type', 'EBITDA Multiple', 'At $5M EBITDA', 'Company Value']
    for i, h in enumerate(val_headers):
        cell = val_tbl.rows[0].cells[i]
        set_cell_bg(cell, '1B2A3D')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    val_data = [
        ('Pure construction', '4x', '$5M', '$20M'),
        ('Mixed (some ITM)', '5.5x', '$5M', '$27.5M'),
        ('Strong ITM (20%+)', '7x', '$5M', '$35M'),
        ('High ITM (40%+)', '9x', '$5M', '$45M'),
    ]
    for i, row_data in enumerate(val_data):
        row = val_tbl.rows[i + 1]
        is_last = i == len(val_data) - 1
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            if is_last:
                set_cell_bg(cell, 'C9A84C')
            elif i % 2 == 0:
                set_cell_bg(cell, 'F5F0E8')
            run = cell.paragraphs[0].add_run(val)
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.bold = is_last
            if is_last:
                run.font.color.rgb = NAVY

    doc.add_paragraph()

    # Woody's 18% table
    add_body(doc, 'Woody\'s 18% — Year 5 Exit Comparison', bold=True, size=11, color=NAVY)
    woody_tbl = doc.add_table(rows=4, cols=6)
    woody_tbl.style = 'Table Grid'
    woody_headers = ['Scenario', 'Year 5 Revenue', 'EBITDA', 'Multiple', 'Company Value', "Woody's 18%"]
    for i, h in enumerate(woody_headers):
        cell = woody_tbl.rows[0].cells[i]
        set_cell_bg(cell, '1B2A3D')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    woody_data = [
        ('Conservative', '$28M', '$4.5M', '5x', '$22.5M', '$4.0M'),
        ('Base Case', '$38M', '$6.5M', '6x', '$39M', '$7.0M'),
        ('Reach', '$50M', '$9M', '7x', '$63M', '$11.3M'),
    ]
    for i, row_data in enumerate(woody_data):
        row = woody_tbl.rows[i + 1]
        is_last = i == len(woody_data) - 1
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            if is_last:
                set_cell_bg(cell, 'C9A84C')
            elif i % 2 == 0:
                set_cell_bg(cell, 'F5F0E8')
            run = cell.paragraphs[0].add_run(val)
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.bold = is_last
            if is_last:
                run.font.color.rgb = NAVY

    doc.add_paragraph()
    add_body(doc,
             "Woody's 18% in the reach scenario = $11.3 million at Year 5 exit. That's on top of the $8.2M cash at close + $2.3M in seller note payments.\n"
             "Total: $21.8M in the base case.  $29.4M in the reach scenario.",
             bold=True, size=11, color=NAVY)

    # Footer
    add_divider(doc)
    add_body(doc, 'GFS Fire Pros | Viking 90 Group | One Buck Capital — CONFIDENTIAL | Not for Distribution', size=8, color=MID_GRAY, italic=True, center=True)

    doc.save('/home/node/.openclaw/workspace/GFS_SafeGuard_Program.docx')
    print('✓ GFS_SafeGuard_Program.docx saved')


# ============================================================
# DOCUMENT 4: GFS Banker Teaser One-Pager
# ============================================================
def build_gfs_banker_teaser():
    doc = Document()
    set_page_margins(doc, 0.75, 0.75, 1.0, 1.0)

    add_navy_header_bar(doc, 'ACQUISITION FINANCING OPPORTUNITY', 'DFW Fire Protection Company — SBA 7(a) + 504 Request')

    # Prepared by block
    doc.add_paragraph()
    prep_tbl = doc.add_table(rows=1, cols=2)
    prep_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_l = prep_tbl.cell(0, 0)
    cell_r = prep_tbl.cell(0, 1)
    for cell in [cell_l, cell_r]:
        set_cell_bg(cell, 'F5F0E8')
    run_l = cell_l.paragraphs[0].add_run('Prepared by:  Joel Bouck — Viking 90 Group\njb@onebuckcapital.com')
    run_l.font.name = 'Calibri'
    run_l.font.size = Pt(10)
    run_l.font.color.rgb = DARK_GRAY
    run_r = cell_r.paragraphs[0].add_run('Date:  March 2026\nCONFIDENTIAL — For Lending Discussion Only')
    run_r.font.name = 'Calibri'
    run_r.font.size = Pt(10)
    run_r.italic = True
    run_r.font.color.rgb = MID_GRAY
    cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_divider(doc)

    # Company overview
    add_heading(doc, 'THE COMPANY', 2, NAVY)
    add_body(doc,
             'A 40-year-old fire suppression, alarm, and inspection company based in Dallas-Fort Worth, Texas. '
             'One of the established independent fire protection contractors in the DFW market — serving nationally recognized commercial general contractors and property owners. '
             'Industry is code-mandated: inspections and maintenance are legally required annually. Recession-resistant, non-discretionary demand.',
             size=10)

    add_divider(doc)

    # Financial snapshot
    add_heading(doc, 'FINANCIAL SNAPSHOT', 2, NAVY)

    fin_tbl = doc.add_table(rows=8, cols=3)
    fin_tbl.style = 'Table Grid'
    fin_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(['Metric', '2024', '2025']):
        cell = fin_tbl.rows[0].cells[i]
        set_cell_bg(cell, '1B2A3D')
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    fin_data = [
        ('Revenue', '$14.2M', '$19.1M  (+34%)'),
        ('Adjusted EBITDA', '~$1.8M', '$2.94M'),
        ('EBITDA Margin', '—', '15.4%'),
        ('Existing Debt / LOC', '—', '$0  ✓'),
        ('Cash on Balance Sheet', '—', '$1.8M'),
        ('Accounts Receivable', '—', '$3.99M'),
        ('Contracted Backlog', '—', '$12M+'),
    ]

    for i, (metric, y24, y25) in enumerate(fin_data):
        row = fin_tbl.rows[i + 1]
        highlight = metric in ('Existing Debt / LOC',)
        for j, val in enumerate([metric, y24, y25]):
            cell = row.cells[j]
            if highlight and j == 2:
                set_cell_bg(cell, 'C9A84C')
            elif i % 2 == 0:
                set_cell_bg(cell, 'F5F0E8')
            run = cell.paragraphs[0].add_run(val)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.bold = highlight and j > 0

    for i, w in enumerate([2.5, 1.8, 2.2]):
        for cell in fin_tbl.columns[i].cells:
            cell.width = Inches(w)

    doc.add_paragraph()

    # Key point callout
    kp_tbl = doc.add_table(rows=1, cols=1)
    cell = kp_tbl.cell(0, 0)
    set_cell_bg(cell, 'C9A84C')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.1)
    run = p.add_run('★  Key Point: The company entered 2026 with zero debt after paying off a $2.4M line of credit over the prior two years — while growing revenue 34%.')
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY

    add_divider(doc)

    # Two-column section: Deal + Financing side by side
    add_heading(doc, 'THE DEAL  &  FINANCING REQUEST', 2, NAVY)

    two_col = doc.add_table(rows=1, cols=2)
    two_col.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_deal = two_col.cell(0, 0)
    cell_fin = two_col.cell(0, 1)

    # Deal table inside cell
    run_deal_hdr = cell_deal.paragraphs[0].add_run('Transaction Structure')
    run_deal_hdr.bold = True
    run_deal_hdr.font.name = 'Georgia'
    run_deal_hdr.font.size = Pt(11)
    run_deal_hdr.font.color.rgb = NAVY

    deal_rows_data = [
        ('Total Enterprise Value', '$10,000,000'),
        ('Business Price', '$7,500,000'),
        ('Real Estate Price', '$2,500,000'),
        ('Seller Note', '$1,150,000 (standby yr 1-2)'),
        ('Seller Equity Rollover', '18% of OpCo'),
    ]
    deal_sub = cell_deal.add_table(rows=len(deal_rows_data), cols=2)
    for i, (term, detail) in enumerate(deal_rows_data):
        r = deal_sub.rows[i]
        if i % 2 == 0:
            set_cell_bg(r.cells[0], 'F5F0E8')
            set_cell_bg(r.cells[1], 'F5F0E8')
        r.cells[0].paragraphs[0].add_run(term).font.name = 'Calibri'
        r.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        r.cells[1].paragraphs[0].add_run(detail).font.name = 'Calibri'
        r.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        r.cells[1].paragraphs[0].runs[0].bold = i == 0

    # Financing table inside other cell
    run_fin_hdr = cell_fin.paragraphs[0].add_run('Financing Request')
    run_fin_hdr.bold = True
    run_fin_hdr.font.name = 'Georgia'
    run_fin_hdr.font.size = Pt(11)
    run_fin_hdr.font.color.rgb = NAVY

    fin_rows_data = [
        ('SBA 7(a)', '$5,000,000', 'Business acquisition'),
        ('SBA 504 — Bank Portion', '$1,250,000', 'Real estate (50%)'),
        ('SBA 504 — CDC', '$1,000,000', 'Real estate (40%)'),
        ('Buyer Equity Injection', '$250,000', 'RE equity (10%)'),
    ]
    fin_sub = cell_fin.add_table(rows=len(fin_rows_data) + 1, cols=3)
    for i, h in enumerate(['Facility', 'Amount', 'Purpose']):
        set_cell_bg(fin_sub.rows[0].cells[i], '1B2A3D')
        run = fin_sub.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
    for i, (fac, amt, pur) in enumerate(fin_rows_data):
        r = fin_sub.rows[i + 1]
        if i % 2 == 0:
            for c in r.cells:
                set_cell_bg(c, 'F5F0E8')
        for j, val in enumerate([fac, amt, pur]):
            r.cells[j].paragraphs[0].add_run(val).font.name = 'Calibri'
            r.cells[j].paragraphs[0].runs[0].font.size = Pt(9)

    cell_deal.width = Inches(3.2)
    cell_fin.width = Inches(3.3)

    doc.add_paragraph()

    # DSCR callout
    dscr_tbl = doc.add_table(rows=1, cols=2)
    dscr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate([
        ('Projected DSCR (Standby Period)', '2.97x'),
        ('Projected DSCR (Full Service)', '2.32x'),
    ]):
        cell = dscr_tbl.cell(0, i)
        set_cell_bg(cell, '1B2A3D')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run_l = p.add_run(f'{label}\n')
        run_l.font.name = 'Calibri'
        run_l.font.size = Pt(9)
        run_l.font.color.rgb = GOLD
        run_v = p.add_run(val)
        run_v.bold = True
        run_v.font.name = 'Georgia'
        run_v.font.size = Pt(18)
        run_v.font.color.rgb = WHITE

    add_divider(doc)

    # Top customers
    add_heading(doc, 'TOP CUSTOMERS', 2, NAVY)
    add_body(doc, 'All investment-grade national general contractors with 10+ year relationships:', size=10, italic=True, color=MID_GRAY)

    cust_tbl = doc.add_table(rows=2, cols=3)
    cust_tbl.style = 'Table Grid'
    cust_data = [
        ('Balfour Beatty Construction', '27% of AR', '10+ year relationship'),
        ('Rogers-O\'Brien Construction', '17% of AR', '10+ year relationship'),
    ]
    for i, (name, pct, note) in enumerate(cust_data):
        for j, val in enumerate([name, pct, note]):
            cell = cust_tbl.rows[i].cells[j]
            if i % 2 == 0:
                set_cell_bg(cell, 'F5F0E8')
            run = cell.paragraphs[0].add_run(val)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.bold = j == 0

    add_body(doc, 'Also: JE Dunn Construction · Structure Tone · Holder Construction Group', size=9, italic=True, color=MID_GRAY)

    add_divider(doc)

    # Buyer background
    add_heading(doc, 'THE BUYER', 2, NAVY)

    buyer_tbl = doc.add_table(rows=1, cols=2)
    buyer_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_bio = buyer_tbl.cell(0, 0)
    cell_stats = buyer_tbl.cell(0, 1)

    set_cell_bg(cell_bio, 'F5F0E8')
    set_cell_bg(cell_stats, '1B2A3D')

    p_bio = cell_bio.paragraphs[0]
    p_bio.paragraph_format.space_before = Pt(6)
    run_bio = p_bio.add_run('Joel Bouck — Viking 90 Group / One Buck Capital')
    run_bio.bold = True
    run_bio.font.name = 'Georgia'
    run_bio.font.size = Pt(11)
    run_bio.font.color.rgb = NAVY

    bio_lines = [
        'Co-founder of Payix, an omni-channel payment technology platform.',
        'Sold to REPAY Holdings (NASDAQ: RPAY) for up to $115 million in January 2022.',
        'Experienced operator with management team in place: CFO and President of Sales both staying post-close with equity incentives.',
    ]
    for line in bio_lines:
        p_line = cell_bio.add_paragraph()
        run_line = p_line.add_run(f'•  {line}')
        run_line.font.name = 'Calibri'
        run_line.font.size = Pt(10)
        run_line.font.color.rgb = DARK_GRAY
    cell_bio.paragraphs[-1].paragraph_format.space_after = Pt(6)

    p_stats = cell_stats.paragraphs[0]
    p_stats.paragraph_format.space_before = Pt(6)
    run_stats_hdr = p_stats.add_run('Capitalization')
    run_stats_hdr.bold = True
    run_stats_hdr.font.name = 'Georgia'
    run_stats_hdr.font.size = Pt(11)
    run_stats_hdr.font.color.rgb = GOLD

    stat_lines = [
        ('Liquid Assets', '$250,000'),
        ('Investor Group', '+$150,000'),
        ('Total Equity', '$400,000'),
    ]
    for label, val in stat_lines:
        p_s = cell_stats.add_paragraph()
        run_s = p_s.add_run(f'{label}:  ')
        run_s.font.name = 'Calibri'
        run_s.font.size = Pt(10)
        run_s.font.color.rgb = OFFWHITE
        run_v = p_s.add_run(val)
        run_v.bold = True
        run_v.font.name = 'Calibri'
        run_v.font.size = Pt(10)
        run_v.font.color.rgb = WHITE
    cell_stats.paragraphs[-1].paragraph_format.space_after = Pt(6)

    cell_bio.width = Inches(4.0)
    cell_stats.width = Inches(2.5)

    add_divider(doc)

    # Why attractive
    add_heading(doc, 'WHY THIS IS AN ATTRACTIVE SBA OPPORTUNITY', 2, NAVY)

    why_items = [
        ('Zero existing debt', 'No competing liens on any collateral. Clean balance sheet.'),
        ('34% revenue growth + $12M backlog', 'Borrowing base expanding. Strong forward visibility.'),
        ('Investment-grade customer base', 'Code-mandated, non-discretionary payments. Balfour Beatty, Rogers-O\'Brien, JE Dunn.'),
        ('Seller rolling 18% equity', 'Strongest possible continuity signal. Skin in the game.'),
        ('Proven operator with exit history', 'REPAY Holdings acquisition, NASDAQ: RPAY, up to $115M.'),
        ('DSCR of 2.32x at full debt service', 'Strong lender coverage. Free cash flow ~$1.67M/year after full debt service.'),
    ]

    why_tbl = doc.add_table(rows=len(why_items), cols=2)
    why_tbl.style = 'Table Grid'
    for i, (title, detail) in enumerate(why_items):
        cell_t = why_tbl.rows[i].cells[0]
        cell_d = why_tbl.rows[i].cells[1]
        if i % 2 == 0:
            set_cell_bg(cell_t, 'C9A84C')
            set_cell_bg(cell_d, 'F5F0E8')
        else:
            set_cell_bg(cell_t, '1B2A3D')
            set_cell_bg(cell_d, 'FFFFFF')
        run_t = cell_t.paragraphs[0].add_run(title)
        run_t.bold = True
        run_t.font.name = 'Calibri'
        run_t.font.size = Pt(9)
        run_t.font.color.rgb = NAVY if i % 2 == 0 else GOLD
        run_d = cell_d.paragraphs[0].add_run(detail)
        run_d.font.name = 'Calibri'
        run_d.font.size = Pt(9)

        cell_t.width = Inches(2.2)
        cell_d.width = Inches(4.3)

    doc.add_paragraph()

    # NDA / contact footer
    nda_tbl = doc.add_table(rows=1, cols=1)
    cell = nda_tbl.cell(0, 0)
    set_cell_bg(cell, '1B2A3D')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Full financials (3-year P&L, balance sheets, tax returns, WIP schedule, AR aging, customer list) available under NDA or standard bank confidentiality agreement.')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = OFFWHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run('Joel Bouck  |  jb@onebuckcapital.com  |  Viking 90 Group  |  CONFIDENTIAL — For Lending Discussion Only')
    run2.bold = True
    run2.font.name = 'Calibri'
    run2.font.size = Pt(9)
    run2.font.color.rgb = GOLD

    doc.save('/home/node/.openclaw/workspace/GFS_Banker_Teaser.docx')
    print('✓ GFS_Banker_Teaser.docx saved')


# ============================================================
# MAIN
# ============================================================
if __name__ == '__main__':
    print('Building documents...')
    build_nnn_broker_list()
    build_gfs_woody_one_pager()
    build_gfs_safeguard()
    build_gfs_banker_teaser()
    print('\nAll 4 documents complete.')
